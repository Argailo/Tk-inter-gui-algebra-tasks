import numpy as np
from task1 import *
from task2 import *
from task3 import *
from task4 import *
from task5 import *
from docx import Document
from docx.shared import Pt
from docx.shared import Length

def quick_cast():
    res1 = gen_task_1((3,1), (1,3), (1, 9))
    res2 = gen_task_2((3, 3), (1, 3), 2)
    res3 = gen_task_3((3, 3), (1, 5))
    res4 = gen_task_4((3, 3), (3, 3), (3, 3), (1, 9))
    res5 = gen_task_5((3, 3), (1, 9), 3, 3)


    # latex_formula_1 = '$\\begin{pmatrix} x_{11}&x_{12} \\\\  x_{21}&x_{22} \end{pmatrix} $'
    latex_formula_1 = res1[3] + res1[4] + res1[5]
    latex_formula_2 = res2[0] + res2[1]
    latex_formula_3 = res3[0] + res3[1] + res3[2]
    latex_formula_4 = res4[0] + res4[1] + res4[2] + res4[3]
    latex_formula_5 = res5[0] + res5[1]

    document = Document()

    document.add_heading('Линейная алгебра, 1 семестр', level=1)

    document.add_heading('1. Перемножение двух матриц', level=2)
    document.add_paragraph(latex_formula_1)
    document.add_heading('2. СЛАУ методом Крамера', level=2)
    document.add_paragraph(latex_formula_2)
    document.add_heading('3. СЛАУ методом обратной матрицы', level=2)
    document.add_paragraph(latex_formula_3)
    document.add_heading('4. Матричные уравнения вида AxB=C', level=2)
    document.add_paragraph(latex_formula_4)
    document.add_heading('5. СЛАУ методом Гаусса', level=2)
    document.add_paragraph(latex_formula_5)

    document.save('Линейная алгебра_1 семестр.docx')