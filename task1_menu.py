import tkinter as tk
from tkinter import ttk
import numpy as np
from task1 import *
from task2 import *
from task3 import *
from task4 import *
from task5 import *
from docx import Document
from docx.shared import Pt
from docx.shared import Length

def create_matrix_entry(parent, rows, cols):
    """Создает сетку Entry для матрицы."""
    matrix_frame = ttk.Frame(parent)
    matrix_frame.grid(row=0, column=0, padx=5, pady=5, sticky="nsew")

    for i in range(rows):
        for j in range(cols):
            entry = ttk.Entry(matrix_frame, width=5)
            entry.grid(row=i, column=j, padx=2, pady=2)

    return matrix_frame


def draw_matrix(canvas, root, matrix):
    rows, cols = matrix.shape
    #canvas = tk.Canvas(root, width=200, height=200)
    #canvas.pack(anchor='n')

    def draw():
        cell_width = canvas.winfo_width() // cols*1
        cell_height = canvas.winfo_height() // rows*1

        for i in range(rows):
            for j in range(cols):
                x1 = j * cell_width
                y1 = i * cell_height
                x2 = x1 + cell_width
                y2 = y1 + cell_height
                canvas.create_rectangle(x1, y1, x2, y2)
                canvas.create_text((x1 + x2) // 2, (y1 + y2) // 2, text=str(matrix[i, j]))

        for i in range(1, rows):
            canvas.create_line(0, i * cell_height, canvas.winfo_width(), i * cell_height)
        for j in range(1, cols):
            canvas.create_line(j * cell_width, 0, j * cell_width, canvas.winfo_height())


    root.after(100, draw) # Задержка в 100 миллисекунд для отрисовки окна
    #root.mainloop()
def God_save(root,res):
    document = Document()
    latex_formula_1 = res[3] + res[4] + res[5]
    document.add_heading('1. Перемножение двух матриц', level=2)
    document.add_paragraph(latex_formula_1)
    document.save('Линейная алгебра_1 семестр.docx')
    root.destroy()
    
    
def first_task_menu_up(root,A1,A2,B1,B2,R1,R2):
    print( (int(A1.get()) , int(A2.get()) ) )
    A_size = (int(A1.get()),int(A2.get()))
    B_size = (int(B1.get()),int(B2.get()))
    range = (int(R1.get()),int(R2.get()))
    root.destroy()
    first_task_menu(A_size,B_size,range)
    
def first_task_menu(A_size, B_size, range):
    root = tk.Tk()
    root.title("Перемножение двух матриц")
    
    style = ttk.Style()
    style.configure("TFrame", background="#bcbcbc")  # Light blue background
    
    # Настройки размеров окна (можно изменить)
    root.geometry("1050x600")
    root.columnconfigure(0, weight=1)  # Растягивает окно по горизонтали
    root.rowconfigure(0, weight=1)     # Растягивает окно по вертикали
    
    
    
    # Верхняя панель с настройками
    top_frame = tk.Frame(root, bg="#40C4FF", height=30)
    top_frame.pack(fill=tk.X)
    
    # Размер матрицы А
    tk.Label(top_frame, text="Размер матрицы А").grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
    tk.Label(top_frame, text="строк").grid(row=1, column=0, sticky=tk.W)
    A_size_1 = tk.Entry(top_frame, width=5)
    print(A_size_1)
    A_size_1.grid(row=1, column=0, sticky=tk.W, padx = 40)
    tk.Label(top_frame, text="столбцов").grid(row=1, column=0, sticky=tk.W, padx =80)
    A_size_2 = tk.Entry(top_frame, width=5)
    A_size_2.grid(row=1, column=0, padx= 140, sticky=tk.W)
    
    # Размер матрицы В
    tk.Label(top_frame, text="Размер матрицы В").grid(row=0, column=1, padx=0, pady=5, sticky=tk.W)
    tk.Label(top_frame, text="строк").grid(row=1, column=1, sticky=tk.W)
    B_size_1 = tk.Entry(top_frame, width=5)
    B_size_1.grid(row=1, column=1, sticky=tk.W, padx = 40)
    tk.Label(top_frame, text="столбцов").grid(row=1, column=1, sticky=tk.W, padx=80)
    B_size_2 = tk.Entry(top_frame, width=5)
    B_size_2.grid(row=1, column=1, sticky=tk.W,padx=140)
    
    # Диапазон значений
    tk.Label(top_frame, text="Диапазон значений коэффициентов").grid(row=0, column=2, padx=0, pady=5,sticky=tk.W)
    tk.Label(top_frame, text="от").grid(row=1, column=2, sticky=tk.W)
    Rg_1 = tk.Entry(top_frame, width=5)
    Rg_1.grid(row=1, column=2, sticky=tk.W, padx = 40)
    tk.Label(top_frame, text="до").grid(row=1, column=2, sticky=tk.W, padx = 80)
    Rg_2 = tk.Entry(top_frame, width=5)
    Rg_2.grid(row=1, column=2, sticky=tk.W, padx = 140)

    print(A_size_1)

    
    # Кнопка обновления и информация
    Update_button = ttk.Button(top_frame, text="↻", command = lambda:first_task_menu_up(root,A_size_1,A_size_2,B_size_1,B_size_2,Rg_1,Rg_2) )
    Update_button.grid(row=1, column=3)
#   generate_button = ttk.Button(root, text="Генерация билетов",command=lambda: gen(data),width=20) 
    ttk.Label(top_frame, text="ⓘ", foreground="red").grid(row=1, column=7)
    
    # Центральная часть с матрицами
    center_frame = ttk.Frame(root)
    center_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
    center_frame.columnconfigure(0, weight=1)
    center_frame.columnconfigure(2, weight=1)
    center_frame.columnconfigure(4, weight=1)
    
    
    canvas1 = tk.Canvas(center_frame, width=200, height=100, bg="#eaeaea")
    canvas1.pack(side=tk.LEFT, padx=10)  # Используем pack для размещения слева
    
    ttk.Label(center_frame, text="X", foreground="black", font = 70).place(x = 225, y=250)
    
    canvas2 = tk.Canvas(center_frame, width=200, height=100, bg="#eaeaea")
    canvas2.pack(side=tk.LEFT, padx=30)  # Используем pack для размещения справа
    
    ttk.Label(center_frame, text="=", foreground="black", font = 130).place(x = 482, y=250)
    
    canvas3 = tk.Canvas(center_frame, width=200, height=100, bg="#eaeaea")
    canvas3.pack(side=tk.LEFT, padx=30)  # Используем pack для размещения справа
    
    preres = gen_task_1(A_size,B_size,range)
    matrix = (preres[0])
    matrix2 = preres[1]
    matrix3 = preres[2]
    
    draw_matrix(canvas1, root, matrix)
    draw_matrix(canvas2, root, matrix2)
    draw_matrix(canvas3, root, matrix3)
    
    
    #Знак умножения
    ttk.Label(center_frame, text="X", font=("Arial", 20))
    
    
    # Знак равно
    #ttk.Label(center_frame, text="=", font=("Arial", 20)).grid(row=0, column=3, pady=20)
    
    # Результирующая матрица
    #matrix3 = create_matrix_entry(center_frame, 3, 2)
    #matrix3.grid(row=0, column=4, padx=10, pady=20)
    
    
    # Нижняя панель с кнопками
    bottom_frame = tk.Frame(root, background="#838383")
    bottom_frame.pack(side=tk.BOTTOM, fill=tk.X)
    Save_button = ttk.Button(bottom_frame, text="Сохранить задание",command= lambda: God_save(root, preres))
    Save_button.grid(row=0, column=0, padx=10, pady=5)
    ttk.Button(bottom_frame, text="Добавить в билет").grid(row=0, column=1, padx=10, pady=5)
    ttk.Button(bottom_frame, text="Следующее задание", image=None).grid(row=0, column=2, padx=10, pady=5)
    
    
    root.mainloop()
